# word_export.py — FullText Safe Split Version (Ver 2.0)
from docx import Document

def export_to_word(sections: dict, save_path: str):
    """
    改良点：
      - FullText などの巨大テキストは 2500 文字ごとに分割して複数段落で出力
      - python-docx の 1段落上限による silent-fail を完全回避
    """

    def safe_add_paragraph(doc, text, chunk_size=2500):
        """
        2500字単位で分割して doc に複数段落で追加する。
        """
        if not text:
            return
        text = text.strip()
        for i in range(0, len(text), chunk_size):
            part = text[i:i+chunk_size]
            doc.add_paragraph(part)

    doc = Document()

    # ---- タイトル（必須）----
    title = sections.get("__TITLE__", "").strip()
    if title:
        doc.add_heading(title, level=0)

    # ---- 本文セクション ----
    for sec, content in sections.items():
        if sec == "__TITLE__":
            continue

        doc.add_heading(sec, level=1)

        # FullText や巨大セクションは安全分割
        safe_add_paragraph(doc, content)

    # ---- 保存 ----
    doc.save(save_path)
